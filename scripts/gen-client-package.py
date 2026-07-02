# -*- coding: utf-8 -*-
"""Generates the NOON. Sprachdienst consolidated client handover package (.docx)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "handover")
DATE = "2 July 2026"
VERSION = "1.0"
BRAND = RGBColor(0x0B, 0x1F, 0x3A)
ACCENT = RGBColor(0xC0, 0x3A, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)

LIVE_URL = "https://www.noon-sprachdienst.de"
ADMIN_URL = "https://www.noon-sprachdienst.de/admin"
VERCEL_URL = "https://vercel.com/noon-sprachdiensts-projects/noon"
FIREBASE_URL = "https://console.firebase.google.com/project/noon-sprachdienst"
FIRESTORE_URL = ("https://console.firebase.google.com/project/noon-sprachdienst/"
                 "firestore/databases/-default-/data/~2FanalyticsEvents")
GITHUB_URL = "https://github.com/noon-sprachdienst/NOON"


def new_doc(title, subtitle):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    brand = doc.add_paragraph()
    r = brand.add_run("NOON. Sprachdienst")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND

    t = doc.add_paragraph()
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = BRAND

    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = GREY

    meta = doc.add_paragraph()
    mr = meta.add_run(f"Client handover package  ·  Version {VERSION}  ·  {DATE}")
    mr.font.size = Pt(9)
    mr.font.color.rgb = GREY
    doc.add_paragraph()
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BRAND
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Number")


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return tbl


def save(doc, filename):
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.abspath(os.path.join(OUT_DIR, filename))
    doc.save(path)
    print("wrote", path)


# ---------------------------------------------------------------------------
# 00 — Start Here / Overview
# ---------------------------------------------------------------------------
def doc_00():
    d = new_doc("Start Here — Project Handover",
                "Overview of the delivered website and your key access points")
    para(d, "This package hands over the complete NOON. Sprachdienst website. It is written for the "
            "business owner: it explains what was delivered, where everything lives, and how to keep "
            "it running. Four further documents accompany this one.")

    h1(d, "The website at a glance")
    table(d, ["Item", "Detail"], [
        ["Live website", LIVE_URL],
        ["Languages", "7 — German (master), English, Arabic (right-to-left), Turkish, Russian, French, Ukrainian"],
        ["Pages", "Homepage, services, pricing, 25 city pages, quote and appointment forms, legal pages"],
        ["Technology", "React 18 + Vite single-page app, hosted on Vercel"],
        ["Contact form", "Sends quote requests by email to info@noon-sprachdienst.de"],
        ["Analytics", "Consent-based, anonymous, stored in Google Firebase (Firestore)"],
        ["Admin dashboard", ADMIN_URL],
    ])

    h1(d, "Your three key access points")
    table(d, ["System", "What it is", "Link"], [
        ["Hosting (Vercel)", "Where the site is built and served from", VERCEL_URL],
        ["Database (Firebase)", "Where anonymous analytics visits are stored", FIREBASE_URL],
        ["Source code (GitHub)", "The complete website source code", GITHUB_URL],
    ])

    h1(d, "What is in this package")
    table(d, ["Document", "Purpose"], [
        ["00 — Start Here", "This overview"],
        ["01 — Access & Accounts", "Every account the site depends on and how to take ownership"],
        ["02 — Deployment & Maintenance", "How updates go live and where to change content"],
        ["03 — SEO & Multilingual Structure", "How the site is found on Google and how the languages work"],
        ["04 — Admin Dashboard & Data Protection", "Using the analytics dashboard and the DSGVO picture"],
    ])

    h1(d, "Business details on record")
    table(d, ["Field", "Value"], [
        ["Company", "NOON. Sprachdienst"],
        ["Headquarters", "Paul-Oeser-Straße 1, 49074 Osnabrück"],
        ["Email", "info@noon-sprachdienst.de"],
        ["Phone", "+49 160 95627666"],
        ["Founded", "2019"],
    ])
    save(d, "NOON_Client_Package_00_Start_Here.docx")


# ---------------------------------------------------------------------------
# 01 — Access & Accounts
# ---------------------------------------------------------------------------
def doc_01():
    d = new_doc("Access & Accounts",
                "The accounts that power the website and how to control them")
    para(d, "The website relies on a small number of external accounts. As the owner you should have "
            "administrator access to each one so the business is never locked out.")

    h1(d, "Accounts and links")
    table(d, ["System", "What it controls", "Link"], [
        ["Vercel", "Hosting, automatic builds, serverless functions, environment variables", VERCEL_URL],
        ["Firebase / Google Cloud", "Firestore database holding anonymous analytics", FIREBASE_URL],
        ["GitHub", "The complete website source code and its history", GITHUB_URL],
        ["Domain (noon-sprachdienst.de)", "The web address and its DNS records", "Managed at the domain registrar"],
        ["Google Workspace (info@)", "The info@noon-sprachdienst.de mailbox and SMTP sending", "Google Admin console"],
    ])

    h1(d, "The analytics database")
    para(d, "Analytics events are stored in a single Firestore collection called analyticsEvents. "
            "You can browse them directly here:")
    bullets(d, [FIRESTORE_URL])
    para(d, "Each record is anonymous and automatically deletes itself 90 days after it is created.")

    h1(d, "Configuration secrets")
    para(d, "Passwords and keys are never stored in the source code. They live only in Vercel under "
            "Project Settings → Environment Variables, and cover: analytics database access, the "
            "admin dashboard password and session key, and the mailbox credentials for the contact "
            "form. Keeping a secure offline copy of these values (for example in a password manager) "
            "means the site can always be rebuilt from scratch.")

    h1(d, "Keeping access secure")
    bullets(d, [
        "Ensure each account above is owned by a business-controlled email, not a personal one.",
        "Ensure billing for the domain, Vercel, and Google is on a business payment method.",
        "Rotate the admin password and the mailbox app password periodically.",
    ])
    save(d, "NOON_Client_Package_01_Access_and_Accounts.docx")


# ---------------------------------------------------------------------------
# 02 — Deployment & Maintenance
# ---------------------------------------------------------------------------
def doc_02():
    d = new_doc("Deployment & Maintenance",
                "How updates go live and where the content lives")
    h1(d, "How hosting works")
    table(d, ["Item", "Value"], [
        ["Framework", "React 18 + Vite (single-page app)"],
        ["Hosting", "Vercel (static hosting plus serverless functions)"],
        ["Source of truth", f"The GitHub repository: {GITHUB_URL}"],
        ["Build command", "npm run build"],
        ["Output folder", "dist"],
        ["Secure connection", "TLS certificate issued and renewed automatically by Vercel"],
    ])

    h1(d, "How an update goes live")
    para(d, "The site deploys automatically. When a change is saved to the main branch of the GitHub "
            "repository, Vercel builds the site and publishes it within a couple of minutes.")
    numbered(d, [
        "A change is committed to the main branch on GitHub.",
        "Vercel detects the change and runs the build.",
        "The new version replaces the live site automatically.",
        "The live website and the contact form are checked once the deployment finishes.",
    ])

    h1(d, "Undoing a change (rollback)")
    para(d, "Every past version is retained by Vercel and can be restored instantly, with no rebuild.")
    numbered(d, [
        f"Open the Vercel project: {VERCEL_URL}",
        "Go to Deployments and find the last version that worked correctly.",
        "Use its menu and choose Promote to Production.",
    ])

    h1(d, "Where the content lives")
    para(d, "All text is defined in the source code. After any change the site is rebuilt and "
            "redeployed so the search-optimised pages regenerate.")
    table(d, ["To change…", "Location in the code"], [
        ["Service and city page text, titles, FAQs", "src/data/seoPages.js"],
        ["Company details, office locations, phone, founding year", "src/data/seoPages.js"],
        ["Buttons, menus and interface labels", "src/data/translations.js"],
        ["Legal texts (Impressum, Datenschutz, AGB)", "src/components/LegalModal.jsx"],
    ])

    h1(d, "Good practice")
    bullets(d, [
        "German is the master language; edit the German text first.",
        "When adding new text, provide the six translations so other languages do not fall back to German.",
        "Keep the legal texts lawful and current; have a lawyer review them when business details change.",
    ])
    save(d, "NOON_Client_Package_02_Deployment_and_Maintenance.docx")


# ---------------------------------------------------------------------------
# 03 — SEO & Multilingual Structure
# ---------------------------------------------------------------------------
def doc_03():
    d = new_doc("SEO & Multilingual Structure",
                "How the site is found on Google and how the languages fit together")
    para(d, "The website is built to be found on search engines. Every page is delivered as ready-made "
            "HTML with its own title, description, and structured data, so Google can read it fully.")

    h1(d, "Languages")
    bullets(d, [
        "Seven languages: German (master), English, Arabic, Turkish, Russian, French, Ukrainian.",
        "Arabic is presented right-to-left.",
        "Each page links to its equivalents in the other languages (hreflang), so Google shows the "
        "right language to each visitor and does not treat them as duplicates.",
        "The language switcher in the top menu moves between the same page in different languages.",
    ])

    h1(d, "City pages")
    para(d, "The site covers 25 cities in total:")
    bullets(d, [
        "Six are physical offices with a real address and opening hours: Osnabrück (headquarters), "
        "Stuttgart, Berlin, Bielefeld, Mainz, and Kiel.",
        "Nineteen are service-area cities that NOON serves digitally across Germany — for example "
        "Hamburg, Köln, München, Frankfurt am Main, and Dresden.",
        "Service-area pages carry no invented address; their structured data honestly states the city "
        "as an area served. This is both accurate and safe for search rankings.",
        "Every city page exists in all seven languages.",
    ])

    h1(d, "Old addresses still work")
    para(d, "The previous website's addresses are preserved. Old links redirect permanently to the "
            "matching new page, so existing search results and bookmarks continue to work instead of "
            "showing an error.")

    h1(d, "Sitemap and search setup")
    table(d, ["Feature", "Status"], [
        ["Ready-made HTML for every page", "Enabled"],
        ["hreflang alternates for 7 languages", "Enabled"],
        ["Structured data (business, services, FAQ, breadcrumb)", "Enabled"],
        ["Automatic sitemap.xml", f"Enabled — {LIVE_URL}/sitemap.xml"],
        ["Unknown addresses shown as a proper not-found page", "Enabled"],
    ])

    h1(d, "Keeping search visibility healthy")
    bullets(d, [
        "Keep the sitemap submitted in Google Search Console; resubmit it after major additions.",
        "After a content change, the site is rebuilt so the pages and sitemap regenerate.",
        "Keep the business name, address, and phone consistent everywhere they appear.",
    ])
    save(d, "NOON_Client_Package_03_SEO_and_Multilingual.docx")


# ---------------------------------------------------------------------------
# 04 — Admin Dashboard & Data Protection
# ---------------------------------------------------------------------------
def doc_04():
    d = new_doc("Admin Dashboard & Data Protection",
                "The analytics dashboard and how visitor data is handled")
    h1(d, "The admin dashboard")
    table(d, ["Item", "Value"], [
        ["Address", ADMIN_URL],
        ["Login", "A single password (kept in Vercel, not in the code)"],
        ["Security", "Signed, secure, HttpOnly session cookie with brute-force protection"],
        ["Shows", "Visits, page views, call-to-action clicks, language, device, browser, and coarse location"],
        ["Data window", "Up to the last 90 days"],
    ])
    para(d, "The dashboard reads from the Firestore analytics collection, which you can also browse "
            "directly:")
    bullets(d, [FIRESTORE_URL])

    h1(d, "Consent and data protection (DSGVO)")
    bullets(d, [
        "Analytics start only after a visitor actively accepts analytics cookies in the consent banner.",
        "If a visitor declines, no analytics data is sent.",
        "No advertising or third-party tracking cookies are used.",
        "No names, emails, IP addresses, or precise coordinates are stored in analytics.",
        "Every record deletes itself automatically 90 days after it is created.",
    ])

    h1(d, "What is recorded per visit")
    table(d, ["Field", "Meaning", "Personal?"], [
        ["Session id", "A random, anonymous identifier", "No"],
        ["Event type", "Page view, page leave, or button click", "No"],
        ["Page and referrer", "Which page was seen and where the visitor came from", "No"],
        ["Language", "Site and browser language", "No"],
        ["Device / browser / OS", "General device and software category", "No"],
        ["Country / city", "Approximate location from the hosting edge", "Coarse only"],
        ["Duration", "Time spent on the page", "No"],
    ])

    h1(d, "Contact-form submissions")
    bullets(d, [
        "Quote and contact requests are delivered by email to info@noon-sprachdienst.de.",
        "They are not stored in any database — they live only in the mailbox.",
        "Attachments are limited to a few files and a small total size, in common document formats.",
    ])
    save(d, "NOON_Client_Package_04_Admin_and_Data_Protection.docx")


if __name__ == "__main__":
    doc_00(); doc_01(); doc_02(); doc_03(); doc_04()
    print("\nClient handover package generated.")
