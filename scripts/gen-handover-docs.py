# -*- coding: utf-8 -*-
"""Generates the NOON. Sprachdienst client handover document set (.docx)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "handover")
DATE = "1 July 2026"
VERSION = "1.0"
BRAND = RGBColor(0x0B, 0x1F, 0x3A)
ACCENT = RGBColor(0xC0, 0x3A, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)


def new_doc(title, subtitle):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    brand = doc.add_paragraph()
    r = brand.add_run("NOON. Sprachdienst")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND

    t = doc.add_paragraph()
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = BRAND

    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = GREY

    meta = doc.add_paragraph()
    mr = meta.add_run(f"Client handover document  ·  Version {VERSION}  ·  {DATE}")
    mr.font.size = Pt(9)
    mr.font.color.rgb = GREY
    doc.add_paragraph()
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BRAND
    p.space_before = Pt(10)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Number")


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)
    return p


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return tbl


def save(doc, filename):
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.abspath(os.path.join(OUT_DIR, filename))
    doc.save(path)
    print("wrote", path)


# ---------------------------------------------------------------------------
# 01 — Credentials & Access Transfer Checklist
# ---------------------------------------------------------------------------
def doc_01():
    d = new_doc("Credentials & Access Transfer Checklist",
                "Accounts and access that must be transferred to the client")
    para(d, "This checklist lists every external account the website depends on. Ownership of each "
            "must be transferred to the client (or a client-controlled email) so the business is not "
            "locked out. Tick each item once the client is the verified owner/admin.")
    h1(d, "Accounts to transfer")
    table(d, ["System", "What it controls", "How to transfer", "Done"], [
        ["Domain registrar (noon-sprachdienst.de)", "The domain name and DNS records",
         "Move to client registrar account or add client as owner; hand over DNS control", "☐"],
        ["Vercel project", "Hosting, builds, serverless functions, environment variables",
         "Transfer project to client Vercel team, or invite client as Owner", "☐"],
        ["Firebase / Google Cloud project", "Firestore analytics database",
         "Add client Google account as Project Owner in IAM; remove developer later", "☐"],
        ["Google Workspace / Gmail (info@)", "The info@noon-sprachdienst.de mailbox",
         "Client controls the Workspace admin and the mailbox password", "☐"],
        ["SMTP app password", "Contact-form email delivery",
         "Regenerate under the client-owned Google account (see Env Var Reference)", "☐"],
        ["GitHub / source repository", "The website source code",
         "Transfer repo ownership or add client as admin", "☐"],
    ])
    h1(d, "Verification steps")
    numbered(d, [
        "Confirm the client can log in to each system above without the developer present.",
        "Confirm billing for the domain, Vercel, and Google is on a client payment method.",
        "Rotate all shared secrets after transfer (admin password, session secret, SMTP password).",
        "Remove the developer's personal access once the client confirms everything works.",
    ])
    note(d, "Keep this document itself confidential — it maps out every access point to the business.")
    save(d, "NOON_01_Credentials_Access_Transfer.docx")


# ---------------------------------------------------------------------------
# 02 — Environment Variables Reference
# ---------------------------------------------------------------------------
def doc_02():
    d = new_doc("Environment Variables Reference",
                "Every configuration secret, where it lives, and how to rotate it")
    para(d, "All secrets are configured in Vercel under Project Settings → Environment Variables "
            "(Production, Preview, and Development scopes). They are never committed to the code; "
            "the repository only contains .env.example as a template.")
    h1(d, "Variables")
    table(d, ["Variable", "Purpose", "Example / format"], [
        ["FIREBASE_PROJECT_ID", "Firebase project identifier for analytics", "noon-analytics"],
        ["FIREBASE_CLIENT_EMAIL", "Service-account email for Firestore access", "…@…​.iam.gserviceaccount.com"],
        ["FIREBASE_PRIVATE_KEY", "Service-account private key", '"-----BEGIN PRIVATE KEY-----\\n…"'],
        ["ADMIN_PASSWORD", "Password for the /admin dashboard login", "long random string"],
        ["ADMIN_SESSION_SECRET", "Signing key for admin session cookies", "separate random string"],
        ["SMTP_HOST", "Outgoing mail server", "smtp.gmail.com"],
        ["SMTP_PORT", "SMTP port", "465"],
        ["SMTP_SECURE", "Use TLS (true unless port 587)", "true"],
        ["SMTP_USER", "Mailbox that sends form emails", "info@noon-sprachdienst.de"],
        ["SMTP_PASSWORD", "Google app password for SMTP_USER", "16-character app password"],
        ["CONTACT_EMAIL", "Where contact-form submissions are delivered", "info@noon-sprachdienst.de"],
    ])
    h1(d, "How to rotate secrets")
    h2(d, "ADMIN_PASSWORD and ADMIN_SESSION_SECRET")
    numbered(d, [
        "Generate two new random values (e.g. run: openssl rand -base64 32).",
        "Update both variables in Vercel → Environment Variables (Production).",
        "Redeploy the project so functions pick up the new values.",
        "Rotating ADMIN_SESSION_SECRET immediately logs out all existing admin sessions.",
    ])
    h2(d, "SMTP_PASSWORD (Google app password)")
    numbered(d, [
        "Sign in to the Google account for info@noon-sprachdienst.de.",
        "Security → 2-Step Verification → App passwords → create a new app password.",
        "Paste it into SMTP_PASSWORD in Vercel and redeploy.",
        "Delete the old app password in the Google account.",
    ])
    h2(d, "FIREBASE_PRIVATE_KEY")
    bullets(d, [
        "In Google Cloud → IAM → Service Accounts, create a new key and update the three FIREBASE_* variables.",
        "Keep the \\n escaped newlines exactly as in .env.example.",
        "Delete the old key after confirming analytics still writes.",
    ])
    note(d, "After changing any variable in Vercel you must trigger a new deployment for it to take effect.")
    save(d, "NOON_02_Environment_Variables_Reference.docx")


# ---------------------------------------------------------------------------
# 03 — Deployment Runbook
# ---------------------------------------------------------------------------
def doc_03():
    d = new_doc("Deployment Runbook",
                "How the site is built, deployed, and rolled back")
    h1(d, "Overview")
    table(d, ["Item", "Value"], [
        ["Framework", "React 18 + Vite (single-page app)"],
        ["Hosting", "Vercel (static hosting + serverless functions)"],
        ["Build command", "npm run build"],
        ["Output directory", "dist"],
        ["Serverless functions", "api/ directory (Node.js)"],
        ["Node version", "20 or newer"],
        ["Routing", "vercel.json rewrites all paths to /index.html (SPA)"],
    ])
    h1(d, "What the build does")
    para(d, "The build command runs two steps:")
    numbered(d, [
        "vite build — compiles the React app into static assets in dist/.",
        "node scripts/prerender-seo.mjs — generates SEO landing pages, localized language "
        "home pages, pricing pages, and sitemap.xml into dist/ so crawlers get real HTML.",
    ])
    h1(d, "Standard deployment (recommended)")
    numbered(d, [
        "Push commits to the production branch (e.g. main) of the connected Git repository.",
        "Vercel automatically builds and deploys.",
        "Verify the live URL loads and the contact form sends a test email.",
    ])
    h1(d, "Manual deployment (fallback)")
    numbered(d, [
        "Install the Vercel CLI: npm i -g vercel.",
        "Run: vercel --prod from the project root.",
    ])
    h1(d, "Rollback")
    numbered(d, [
        "Open the Vercel dashboard → Deployments.",
        "Find the last known-good deployment.",
        "Use the ⋯ menu → Promote to Production (instant rollback, no rebuild).",
    ])
    h1(d, "Domain & DNS")
    bullets(d, [
        "The domain noon-sprachdienst.de is added under Vercel → Settings → Domains.",
        "DNS points to Vercel (A record or CNAME as shown in the Vercel domain panel).",
        "TLS certificates are issued and renewed automatically by Vercel.",
    ])
    note(d, "Environment variables must be set in Vercel before the first production deploy or the "
            "contact form and analytics will return service-unavailable errors.")
    save(d, "NOON_03_Deployment_Runbook.docx")


# ---------------------------------------------------------------------------
# 04 — Admin Dashboard Guide
# ---------------------------------------------------------------------------
def doc_04():
    d = new_doc("Admin Dashboard Guide",
                "Using the protected analytics dashboard")
    h1(d, "Access")
    table(d, ["Item", "Value"], [
        ["URL", "https://noon-sprachdienst.de/admin"],
        ["Login", "Enter the ADMIN_PASSWORD value"],
        ["Session length", "12 hours, then login is required again"],
        ["Session security", "HMAC-signed, HttpOnly, Secure, SameSite=Strict cookie"],
        ["Brute-force protection", "Max 10 login attempts per 15 minutes per IP"],
    ])
    h1(d, "What the dashboard shows")
    para(d, "The dashboard reads anonymous, consent-based visit data from Firestore for up to the "
            "last 90 days. Typical metrics include:")
    bullets(d, [
        "Page views and page-leave events",
        "CTA (call-to-action) clicks",
        "Site language and browser language",
        "Device type (desktop / tablet / mobile), browser, and OS",
        "Approximate country and city (from Vercel edge headers)",
        "Visit duration and referrer",
    ])
    h1(d, "Logging out")
    para(d, "Use the logout action in the dashboard. This clears the session cookie immediately.")
    note(d, "Only aggregate, anonymous data is collected, and only after the visitor accepts analytics "
            "cookies. See the DSGVO / Analytics Data Sheet for the full data-protection picture.")
    save(d, "NOON_04_Admin_Dashboard_Guide.docx")


# ---------------------------------------------------------------------------
# 05 — Content Editing Guide
# ---------------------------------------------------------------------------
def doc_05():
    d = new_doc("Content Editing Guide",
                "Where the website text lives and how to change it")
    para(d, "All content is defined in code files. After any change you must rebuild and redeploy "
            "(see the Deployment Runbook) so the prerendered SEO pages are regenerated.")
    h1(d, "Where to edit what")
    table(d, ["To change…", "Edit this file"], [
        ["SEO landing pages (titles, descriptions, body, FAQs)", "src/data/seoPages.js"],
        ["Language home page meta (all 7 languages)", "src/data/seoPages.js → LANGUAGE_HOME_META"],
        ["Company details, locations, phone, founding year", "src/data/seoPages.js → COMPANY / LOCATIONS"],
        ["Interface strings (buttons, menus, labels)", "src/data/translations.js"],
        ["Service navigation and service sheets", "src/data/serviceContent.js"],
        ["Service sheet body text and translations", "src/content/services/sheets.jsx"],
        ["Legal texts (Impressum, Datenschutz, AGB)", "src/components/LegalModal.jsx"],
    ])
    h1(d, "Golden rules")
    bullets(d, [
        "German is the master language. Add or edit the German text first.",
        "When you add German text to a service sheet, add its translation for en, ar, tr, ru, fr, uk "
        "in the translation map, or it will fall back to German/English on those pages.",
        "For French and Ukrainian, use the typographic apostrophe ’ (not ') inside single-quoted "
        "strings so the code does not break.",
        "Always run npm run build locally before deploying; it fails fast if a file has a syntax error.",
    ])
    h1(d, "Verifying translations")
    para(d, "Two helper scripts check translation health after edits:")
    bullets(d, [
        "npm run smoke:services — renders every service page in every language and fails on crash.",
        "node scripts/check-map-completeness.mjs — lists any service text missing a translation.",
    ])
    save(d, "NOON_05_Content_Editing_Guide.docx")


# ---------------------------------------------------------------------------
# 06 — DSGVO / GDPR & Analytics Data Sheet
# ---------------------------------------------------------------------------
def doc_06():
    d = new_doc("DSGVO / GDPR & Analytics Data Sheet",
                "What visitor data is collected, how, and for how long")
    para(d, "This sheet documents the website's data processing so the client can answer a data-"
            "protection inquiry. It complements ANALYTICS_SETUP.md in the repository.")
    h1(d, "Consent model")
    bullets(d, [
        "Analytics only start after the visitor actively accepts analytics cookies in the consent banner.",
        "If the visitor declines or withdraws consent, no analytics events are sent.",
        "No advertising or third-party tracking cookies are used.",
    ])
    h1(d, "Data collected per event")
    table(d, ["Field", "Meaning", "Personal?"], [
        ["sessionId", "Random anonymous session id (not linked to identity)", "No"],
        ["type", "Event type (page_view, page_leave, cta_click, …)", "No"],
        ["path / referrer", "Page visited and referring domain", "No"],
        ["siteLanguage / browserLanguage", "Language settings", "No"],
        ["device / browser / os", "Device and software category", "No"],
        ["country / city", "Approx. location from Vercel edge headers", "Coarse only"],
        ["durationSeconds", "Time on page", "No"],
        ["createdAt / expireAt", "Timestamp and auto-delete date", "No"],
    ])
    para(d, "No names, emails, IP addresses, or precise coordinates are stored in analytics.")
    h1(d, "Retention")
    bullets(d, [
        "Every analytics record carries an expireAt of 90 days after creation.",
        "Firestore's TTL policy automatically deletes records past expireAt.",
        "The admin dashboard can only query the last 90 days.",
    ])
    h1(d, "Contact-form data")
    bullets(d, [
        "The quote/contact form sends the submission by email to info@noon-sprachdienst.de.",
        "Form data is not stored in a database — it lives only in the mailbox.",
        "Attachments are limited to 6 files / 3 MB total, PDF, Word, or images only.",
    ])
    note(d, "Keep the Datenschutzerklärung (privacy policy) in LegalModal.jsx consistent with this sheet.")
    save(d, "NOON_06_DSGVO_Analytics_Data_Sheet.docx")


# ---------------------------------------------------------------------------
# 07 — Legal Content Source-of-Truth
# ---------------------------------------------------------------------------
def doc_07():
    d = new_doc("Legal Content Source-of-Truth",
                "Where the Impressum, Datenschutz, and AGB text lives")
    para(d, "The legal pages are rendered from code, not a CMS. The client is responsible for keeping "
            "this text lawful and up to date.")
    h1(d, "Location")
    table(d, ["Legal text", "Where it is defined"], [
        ["Impressum (imprint)", "src/components/LegalModal.jsx"],
        ["Datenschutzerklärung (privacy policy)", "src/components/LegalModal.jsx"],
        ["AGB (terms & conditions)", "src/components/LegalModal.jsx"],
    ])
    h1(d, "How to update")
    numbered(d, [
        "Open src/components/LegalModal.jsx.",
        "Edit the relevant section text (German master).",
        "Rebuild and redeploy (see Deployment Runbook).",
        "Have a lawyer review Impressum and Datenschutz whenever business details change.",
    ])
    h1(d, "Business facts referenced in legal text")
    table(d, ["Field", "Value"], [
        ["Company", "NOON. Sprachdienst"],
        ["Headquarters", "Paul-Oeser-Straße 1, 49074 Osnabrück"],
        ["Email", "info@noon-sprachdienst.de"],
        ["Phone", "+49 160 95627666"],
        ["Founded", "2019"],
    ])
    note(d, "These same facts appear in the homepage structured data and the location pages — if the "
            "business address or phone changes, update seoPages.js and index.html as well.")
    save(d, "NOON_07_Legal_Content_Source_of_Truth.docx")


# ---------------------------------------------------------------------------
# 08 — Backup & Recovery
# ---------------------------------------------------------------------------
def doc_08():
    d = new_doc("Backup & Recovery",
                "Protecting the code and the analytics data")
    h1(d, "Website code")
    bullets(d, [
        "The entire site lives in the Git repository — that is the primary backup.",
        "Every Vercel deployment is retained and can be re-promoted, so any past version is recoverable.",
        "Keep at least one copy of the repository outside the developer's machine (GitHub + a clone).",
    ])
    h1(d, "Analytics data (Firestore)")
    para(d, "Analytics is non-critical, anonymous, and auto-expires after 90 days, so backups are "
            "optional. If a snapshot is wanted:")
    numbered(d, [
        "In Google Cloud Console, open Firestore → Import/Export.",
        "Export the analyticsEvents collection to a Cloud Storage bucket.",
        "Download or schedule periodic exports as needed.",
    ])
    h1(d, "Environment variables")
    bullets(d, [
        "Store a secure offline copy of all Vercel environment variable values (a password manager).",
        "Without these, a fresh deployment cannot connect to email or Firestore.",
    ])
    h1(d, "Recovery drill")
    numbered(d, [
        "Clone the repository to a clean machine.",
        "Create a new Vercel project from it.",
        "Restore the environment variables.",
        "Point the domain at the new project.",
    ])
    save(d, "NOON_08_Backup_Recovery.docx")


# ---------------------------------------------------------------------------
# 09 — SEO Maintenance Checklist
# ---------------------------------------------------------------------------
def doc_09():
    d = new_doc("SEO Maintenance Checklist",
                "Keeping search visibility healthy over time")
    para(d, "The site ships with prerendered HTML per page, hreflang alternates for 7 languages, "
            "structured data, robots.txt, and a generated sitemap.xml. This checklist keeps it healthy. "
            "It extends SEO_SEARCH_CONSOLE_CHECKLIST.md in the repository.")
    h1(d, "One-time setup")
    numbered(d, [
        "Verify the domain in Google Search Console (and Bing Webmaster Tools).",
        "Submit https://noon-sprachdienst.de/sitemap.xml in Search Console.",
        "Confirm robots.txt is reachable and references the sitemap.",
    ])
    h1(d, "After every content change")
    bullets(d, [
        "Rebuild and redeploy so prerendered pages and sitemap.xml regenerate.",
        "Use Search Console → URL Inspection → Request Indexing for important changed pages.",
    ])
    h1(d, "Monthly")
    bullets(d, [
        "Review Search Console coverage for errors or newly excluded pages.",
        "Check that the homepage business data (name, address, phone) still matches reality.",
        "Confirm no unexpected 404s are being indexed (they are served with noindex).",
    ])
    h1(d, "Built-in SEO facts to preserve")
    table(d, ["Feature", "Status"], [
        ["Prerendered HTML per service / location / pricing page", "Enabled"],
        ["hreflang + x-default alternates (7 languages)", "Enabled"],
        ["Structured data (Organization, Service, FAQ, Breadcrumb, ProfessionalService)", "Enabled"],
        ["Localized language home pages with real crawlable body", "Enabled"],
        ["Unknown URLs served with noindex", "Enabled"],
    ])
    save(d, "NOON_09_SEO_Maintenance_Checklist.docx")


# ---------------------------------------------------------------------------
# 10 — Architecture / Tech-Stack Overview
# ---------------------------------------------------------------------------
def doc_10():
    d = new_doc("Architecture & Tech-Stack Overview",
                "A one-page mental model of how the site is built")
    h1(d, "Stack")
    table(d, ["Layer", "Technology"], [
        ["Frontend", "React 18, Vite, react-helmet-async"],
        ["Maps & visuals", "Leaflet (branch maps), Three.js + three-globe (hero)"],
        ["Icons", "Font Awesome, Lucide"],
        ["Server endpoints", "Vercel serverless functions (Node.js) in api/"],
        ["Email", "Nodemailer via Google Workspace SMTP"],
        ["Analytics store", "Firebase Admin SDK + Cloud Firestore"],
        ["Hosting", "Vercel (static + functions)"],
    ])
    h1(d, "How routing works")
    bullets(d, [
        "It is a single-page app with a hand-rolled router in src/App.jsx.",
        "vercel.json rewrites all URLs to /index.html; React then renders the matching view.",
        "A build-time script (scripts/prerender-seo.mjs) writes real HTML per SEO page for crawlers.",
        "Seven languages: de (master), en, ar, tr, ru, fr, uk. Arabic renders right-to-left.",
    ])
    h1(d, "Server endpoints")
    table(d, ["Endpoint", "Purpose"], [
        ["POST /api/contact/send", "Contact / quote form → email (honeypot + time-trap protected)"],
        ["POST /api/analytics/event", "Record an anonymous analytics event (rate-limited)"],
        ["POST /api/admin/login", "Admin dashboard login (rate-limited)"],
        ["GET /api/admin/status", "Check current admin session"],
        ["POST /api/admin/logout", "End admin session"],
        ["GET /api/admin/analytics", "Read analytics for the dashboard (auth required)"],
    ])
    h1(d, "Key source folders")
    table(d, ["Path", "Contents"], [
        ["src/components, src/pages", "UI components and page views"],
        ["src/data", "SEO pages, translations, service content, company facts"],
        ["src/content/services", "Service sheet templates and translations"],
        ["api", "Serverless functions and shared _lib helpers"],
        ["scripts", "Build-time prerender and QA scripts"],
    ])
    save(d, "NOON_10_Architecture_Overview.docx")


# ---------------------------------------------------------------------------
# 11 — Third-party Asset Licenses
# ---------------------------------------------------------------------------
def doc_11():
    d = new_doc("Third-party Asset Licenses",
                "Record of external assets and their licensing")
    para(d, "The site uses roughly 200 icon and image assets under public/assets, plus several open-"
            "source libraries. This document records their licensing so the client is protected. "
            "The developer should complete the 'Source & license' column for each asset group before "
            "final sign-off.")
    h1(d, "Image / icon assets")
    table(d, ["Asset group (public/assets/…)", "Source & license", "Attribution required?"], [
        ["Service icon sets (per-service PNGs)", "____________________", "☐"],
        ["Branch / location photos", "____________________", "☐"],
        ["Hero and world-map imagery", "____________________", "☐"],
        ["Logos and brand marks", "Owned by NOON. Sprachdienst", "No"],
    ])
    note(d, "If any icons come from a stock icon pack (e.g. Flaticon), the pack's license and any "
            "required attribution must be documented here or the icons replaced with licensed ones.")
    h1(d, "Open-source libraries")
    table(d, ["Library", "License"], [
        ["React, React DOM", "MIT"],
        ["Vite", "MIT"],
        ["Leaflet", "BSD-2-Clause"],
        ["Three.js, three-globe", "MIT"],
        ["Font Awesome (free)", "Icons CC BY 4.0 / Fonts SIL OFL / Code MIT"],
        ["Lucide", "ISC"],
        ["Nodemailer", "MIT"],
        ["firebase-admin", "Apache-2.0"],
        ["react-helmet-async", "Apache-2.0"],
    ])
    save(d, "NOON_11_Third_Party_Asset_Licenses.docx")


# ---------------------------------------------------------------------------
# 12 — Ownership / Handover Letter
# ---------------------------------------------------------------------------
def doc_12():
    d = new_doc("Ownership & Handover Letter",
                "Confirmation of delivery and transfer of the website")
    para(d, f"Date: {DATE}")
    doc_para = para(d, "")
    para(d, "This letter confirms the delivery of the NOON. Sprachdienst website and the transfer of "
            "all associated code and intellectual property to the client.")
    h1(d, "Delivered")
    bullets(d, [
        "The complete website source code (frontend, serverless functions, build scripts).",
        "A multilingual production site in 7 languages with prerendered SEO pages.",
        "A consent-aware analytics dashboard and an email-based contact/quote form.",
        "This handover document set (12 documents).",
    ])
    h1(d, "Transferred to the client")
    bullets(d, [
        "Ownership of the source code and all custom content created for this project.",
        "Administrative access to hosting, domain, email, and the analytics project "
        "(per the Credentials & Access Transfer Checklist).",
    ])
    h1(d, "Client responsibilities after handover")
    bullets(d, [
        "Keep all credentials secure and rotate the shared secrets.",
        "Keep legal texts (Impressum, Datenschutz, AGB) lawful and current.",
        "Maintain hosting, domain, and email subscriptions.",
    ])
    para(d, "")
    para(d, "Delivered by: ______________________________    Date: ______________")
    para(d, "")
    para(d, "Accepted by (client): ________________________    Date: ______________")
    save(d, "NOON_12_Ownership_Handover_Letter.docx")


if __name__ == "__main__":
    doc_01(); doc_02(); doc_03(); doc_04(); doc_05(); doc_06()
    doc_07(); doc_08(); doc_09(); doc_10(); doc_11(); doc_12()
    print("\nAll handover documents generated.")
